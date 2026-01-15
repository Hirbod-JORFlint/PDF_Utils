#!/usr/bin/env python3
"""
pdf_to_text_docx_high_fidelity.py

High-Fidelity PDF -> DOCX Converter
Improvements over original:
 - Matches DOCX page size to PDF page size.
 - Intelligent Font Name mapping (PDF internal names -> Windows/Word names).
 - Hyperlink detection and embedding.
 - Superscript detection.
 - Zero-spacing layout (prevents large gaps between lines).
 - In-memory image processing (no temp files).
Usage examples:
  python pdf_to_text_docx_with_columns.py input.pdf --out docx --output out.docx
  python pdf_to_text_docx_with_columns.py input.pdf --out txt --output out.txt --ocr
  python pdf_to_text_docx_with_columns.py input.pdf --out docx --latex --max-columns 3 --col-gap 0.18
Usage:
  python pdf_to_text_docx_high_fidelity.py input.pdf --output out.docx
"""

import argparse
import os
import sys
import re
import io

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.shared import OxmlElement, qn as _qn
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import statistics

# ==========================================
# Helpers
# ==========================================

def int_to_rgb(color_value):
    """Convert PyMuPDF color int/tuple to RGB tuple."""
    if isinstance(color_value, int):
        # Some PDFs store as simple int
        r = (color_value >> 16) & 255
        g = (color_value >> 8) & 255
        b = color_value & 255
        return (r, g, b)
    elif isinstance(color_value, (list, tuple)) and len(color_value) >= 3:
        # PyMuPDF floats 0.0-1.0 or ints 0-255
        if all(isinstance(c, float) for c in color_value):
            return (int(color_value[0]*255), int(color_value[1]*255), int(color_value[2]*255))
        return tuple(map(int, color_value[:3]))
    return (0, 0, 0)

def map_font_name(fname, flags=None):
    """
    Map PDF font names to standard Word fonts to improve visual fidelity.
    - Normalizes subset prefixes (AAAAAA+FontName), suffix tokens, and casing.
    - Tries to preserve family (serif/sans/mono) and common names.
    """
    if not fname:
        return "Calibri"

    # Remove subset prefix like "ABCDEE+FontName" and any encoding suffix
    fname = fname.split('+')[-1]
    fname = fname.split('-')[0]

    # Clean common suffix tokens and punctuation
    fname_clean = re.sub(r'^[A-Z]{6}\+', '', fname)
    fname_clean = re.sub(r'[-_,.](bold|italic|regular|mt|ps|std|roman|oblique|semi|condensed|narrow)$',
                        '', fname_clean, flags=re.I).strip()
    fname_lower = fname_clean.lower()

    # family heuristics
    if "times" in fname_lower or "serif" in fname_lower:
        return "Times New Roman"
    if "arial" in fname_lower or "helvetica" in fname_lower:
        return "Arial"
    if "courier" in fname_lower or "mono" in fname_lower or "monospace" in fname_lower:
        return "Courier New"
    if "calibri" in fname_lower:
        return "Calibri"
    if "cambria" in fname_lower:
        return "Cambria"
    if "georgia" in fname_lower:
        return "Georgia"

    # Fallback:
    return fname_clean or "Calibri"

def add_hyperlink(paragraph, url, text, color, is_bold, is_italic, font_name, font_size):
    """
    Manually inject a hyperlink into the DOCX XML.
    python-docx does not natively support adding links to runs easily.
    """
    # This gets the paragraph's relation part ID
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    # Create the <w:hyperlink> tag
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create the <w:r> (run) tag
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Color (use source color)
    c_el = OxmlElement('w:color')
    c_el.set(qn('w:val'), "{:02x}{:02x}{:02x}".format(*color))
    rPr.append(c_el)

    # Underline
    u_el = OxmlElement('w:u')
    u_el.set(qn('w:val'), 'single')
    rPr.append(u_el)

    # Font family (rFonts) so Word renders link in intended face
    if font_name:
        rfonts = OxmlElement('w:rFonts')
        rfonts.set(qn('w:ascii'), font_name)
        rfonts.set(qn('w:hAnsi'), font_name)
        rfonts.set(qn('w:eastAsia'), font_name)
        rPr.append(rfonts)

    # Font size (size is in half-points in docx XML)
    try:
        if font_size:
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), str(int(float(font_size) * 2)))
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), str(int(float(font_size) * 2)))
            rPr.append(szCs)
    except Exception:
        pass

    # Bold/Italic
    if is_bold:
        rPr.append(OxmlElement('w:b'))
    if is_italic:
        rPr.append(OxmlElement('w:i'))
    
    # Text
    text_el = OxmlElement('w:t')
    text_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    text_el.text = text
    
    run.append(rPr)
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)

def is_box_inside(inner, outer):
    ix0, iy0, ix1, iy1 = inner
    ox0, oy0, ox1, oy1 = outer
    cx = (ix0 + ix1) / 2
    cy = (iy0 + iy1) / 2
    return (ox0 <= cx <= ox1) and (oy0 <= cy <= oy1)

# ==========================================
# Extraction Logic
# ==========================================

def get_link_target(bbox, links):
    """Find if a bbox intersects with a known link area."""
    # links is a list of dicts from page.get_links()
    cx = (bbox[0] + bbox[2]) / 2
    cy = (bbox[1] + bbox[3]) / 2
    
    for link in links:
        lb = link['from'] # Rect
        if lb.x0 <= cx <= lb.x1 and lb.y0 <= cy <= lb.y1:
            return link.get('uri')
    return None

def analyze_columns_simple(spans, page_width):
    """
    Improved greedy clustering of element centers + width-awareness.

    - Uses bbox centers (x_c) and element widths to cluster.
    - Threshold scales with page width and element width (adaptive).
    - Returns column center x-coordinates.
    """
    if not spans:
        return [page_width / 2.0]

    elems = []
    for s in spans:
        bbox = s.get('bbox') or (0, 0, 0, 0)
        x0, _, x1, _ = bbox
        center = (x0 + x1) / 2.0
        width = max(1.0, x1 - x0)
        elems.append((center, width))

    # sort by center
    elems.sort(key=lambda e: e[0])

    columns = []
    current = [elems[0][0]]
    # adaptive threshold: smaller for narrow pages, larger for wide
    base_thresh = max(page_width * 0.08, 24)  # 8% of page or 24pt floor

    for center, width in elems[1:]:
        # use width to loosen threshold for wide elements (so they don't split columns)
        adaptive = max(base_thresh, width * 0.6)
        if abs(center - current[-1]) <= adaptive:
            current.append(center)
        else:
            columns.append(sum(current) / len(current))
            current = [center]
    if current:
        columns.append(sum(current) / len(current))

    # merge close column centers to avoid micro-columns
    merged = []
    merge_thresh = max(page_width * 0.03, 18)  # 3% of page or 18pt floor
    for c in columns:
        if not merged:
            merged.append(c)
        else:
            if abs(c - merged[-1]) <= merge_thresh:
                # average into last
                merged[-1] = (merged[-1] + c) / 2.0
            else:
                merged.append(c)

    # Ensure at least one center (fallback to page middle)
    if not merged:
        return [page_width / 2.0]
    return merged

def extract_page_content(page):
    elements = []
    links = page.get_links()

    # 1. Tables
    table_rects = []
    tables = page.find_tables()
    for tab in tables:
        bbox = tab.bbox
        table_rects.append(bbox)
        elements.append({
            'type': 'table',
            'bbox': bbox,
            'y_sort': bbox[1],
            'data': tab.extract()
        })

    # 2. Text & Images
    text_blocks = page.get_text("dict", sort=True)["blocks"]

    page_height = page.rect.height
    header_threshold = page_height * 0.08  # Top 8%
    footer_threshold = page_height * 0.92  # Bottom 8%
    
    for b in text_blocks:
        bbox = b['bbox']
        # Identify if block is a Header/Footer
        is_margin_content = bbox[3] < header_threshold or bbox[1] > footer_threshold
        
        # Optional: Skip common artifacts like single-digit page numbers in margins
        if is_margin_content and len(b.get("lines", [])) == 1:
            if re.match(r'^\d+$', page.get_text("text", clip=bbox).strip()):
                continue
            
        if any(is_box_inside(bbox, tr) for tr in table_rects):
            continue

        if b['type'] == 0:  # Text
            para_lines = []
            for line in b['lines']:
                line_spans = []
                for span in line['spans']:
                    uri = get_link_target(span['bbox'], links)
                    line_spans.append({
                        'text': span['text'],
                        'font': span['font'],
                        'size': span['size'],
                        'color': span['color'],
                        'flags': span['flags'],
                        'bbox': span['bbox'],
                        'origin': span['origin'],
                        'link': uri
                    })
                if line_spans:
                    para_lines.append({'bbox': line['bbox'], 'spans': line_spans})
            
            if para_lines:
                paragraphs = []
                def avg_font_size(line):
                    sizes = [s.get('size') for s in line['spans'] if s.get('size')]
                    return float(sizes[0]) if sizes else 12.0

                current = {'bbox': para_lines[0]['bbox'], 'spans': list(para_lines[0]['spans'])}
                prev_center_y = (para_lines[0]['bbox'][1] + para_lines[0]['bbox'][3]) / 2.0
                current_font = avg_font_size(para_lines[0])

                def representative_font_size(line):
                    sizes = [s.get('size') for s in line['spans'] if s.get('size')]
                    if not sizes:
                        return 12.0
                    # use median (robust to outliers)
                    return float(statistics.median(sizes))
                
                current_font = representative_font_size(para_lines[0])
                prev_bottom = para_lines[0]['bbox'][3]  # y1 of previous bbox

                for ln in para_lines[1:]:
                    top = ln['bbox'][1]
                    # vertical gap: distance from previous bottom to this top (PDF coords)
                    gap_pts = max(0.0, top - prev_bottom)
                    ln_font = representative_font_size(ln)

                    # threshold: small absolute minimum + fraction of larger font size for robustness
                    min_gap_pts = 3.0
                    rel_gap_factor = 0.45
                    threshold = max(min_gap_pts, rel_gap_factor * max(current_font, ln_font))

                    # horizontal alignment: left indent difference
                    x_offset = abs(ln['bbox'][0] - current['bbox'][0])

                    if gap_pts <= threshold and x_offset < 12:
                        # same paragraph: merge. handle hyphenation across lines
                        last_span = current['spans'][-1]
                        if last_span['text'].rstrip().endswith('-') and len(last_span['text'].strip()) > 1:
                            last_span['text'] = last_span['text'].rstrip().rstrip('-')
                        else:
                            if not last_span['text'].endswith(' '):
                                last_span['text'] += ' '
                        current['spans'].extend(ln['spans'])
                        # expand bbox
                        c = list(current['bbox'])
                        c[2] = max(c[2], ln['bbox'][2])
                        c[3] = max(c[3], ln['bbox'][3])
                        current['bbox'] = tuple(c)
                        prev_bottom = current['bbox'][3]
                        # keep current_font as a representative (max of values to avoid shrink)
                        current_font = max(current_font, ln_font)
                    else:
                        # new paragraph
                        paragraphs.append(current)
                        current = {'bbox': ln['bbox'], 'spans': list(ln['spans'])}
                        prev_bottom = current['bbox'][3]
                        current_font = ln_font

                paragraphs.append(current)

                for para in paragraphs:
                    elements.append({
                        'type': 'text',
                        'bbox': para['bbox'],
                        'y_sort': para['bbox'][1],
                        'lines': [para],
                        'indent_pt': para['bbox'][0]
                    })
                    
        elif b['type'] == 1:  # Image
            img_bytes = b.get("image")
            if img_bytes:
                elements.append({
                    'type': 'image',
                    'bbox': bbox,
                    'y_sort': bbox[1],
                    'bytes': img_bytes,
                    'ext': b.get("ext", "png")
                })

    # Detect logical columns based on the distribution of x0 coordinates
    # --- Replace local clustering with centralized function -----------------
    # Build a lightweight 'spans' list (each item has a bbox) compatible
    # with analyze_columns_simple(spans, page_width)
    spans_for_columns = [{'bbox': el['bbox']} for el in elements if el.get('bbox')]
    columns = analyze_columns_simple(spans_for_columns, page.rect.width)

    # assign each element to the nearest column center, then sort by (col, y, x)
    def get_column_index_for_bbox(bbox, columns):
        # assign based on bbox center; if element is very wide and overlaps two columns,
        # choose the column whose center lies inside the bbox, else nearest center.
        x0, _, x1, _ = bbox
        center = (x0 + x1) / 2.0
        # prefer columns inside bbox range
        for idx, c in enumerate(columns):
            # if c lies inside the element bbox, assign to that column
            if x0 <= c <= x1:
                return idx
        # fallback: nearest column center
        best_idx, best_d = 0, abs(center - columns[0])
        for idx, c in enumerate(columns[1:], start=1):
            d = abs(center - c)
            if d < best_d:
                best_idx, best_d = idx, d
        return best_idx

    def get_reading_order(el):
        bx0, by0, bx1, by1 = el['bbox']
        height = max(0.0, by1 - by0)
        col_idx = get_column_index_for_bbox(el['bbox'], columns)
        # within a column: sort top-to-bottom (use by0), prefer taller blocks earlier for tie-breaks,
        # then left-to-right.
        return (col_idx, float(by0), -float(height), float(bx0))

    # --- Attach short centered text blocks immediately below images as captions ---
    # Approx: if a text block's top is within ~18pt below an image bottom and its width is
    # similar to the image, consider it a caption.
    new_elements = list(elements)  # shallow copy for mutation

    image_elements = [el for el in elements if el['type'] == 'image']
    consumed_text_ids = set()
    for img in image_elements:
        ix0, iy0, ix1, iy1 = img['bbox']
        img_mid_x = (ix0 + ix1) / 2.0
        img_w = max(1.0, ix1 - ix0)
        img_h = max(1.0, iy1 - iy0)

        # adaptive thresholds
        max_vdist = max(18.0, img_h * 0.35)   # up to 35% of image height
        max_h_center = max(12.0, img_w * 0.4) # allow a bit more horizontal variance

        for idx, txt in enumerate(elements):
            # skip if it's the image itself or not text or already consumed
            if txt is img or txt.get('type') != 'text' or id(txt) in consumed_text_ids:
                continue
            tx0, ty0, tx1, ty1 = txt['bbox']
            vdist = ty0 - iy1
            h_center_diff = abs(((tx0 + tx1) / 2.0) - img_mid_x)

            # require shortish caption and modest font size
            cap_text = "".join([s['text'] for l in txt.get('lines', []) for s in l.get('spans', [])]).strip()
            if not cap_text:
                continue
            # estimate average char width (pts). Try to read first span size to estimate; fallback to 6.0 pts
            est_font_pt = None
            try:
                first_line = txt.get('lines', [])[0]
                first_span = first_line.get('spans', [])[0]
                if first_span and first_span.get('size'):
                    est_font_pt = float(first_span.get('size'))
            except Exception:
                est_font_pt = None
            if not est_font_pt:
                est_font_pt = 6.0

            # conservative char width estimate: about 0.5 * font_pt (typical for proportional fonts)
            char_width_pts = max(2.5, 0.45 * est_font_pt)
            max_chars = int((img_w * 0.8) / char_width_pts)
            # reject overly long text relative to image width
            if len(cap_text) > max_chars:
                continue
            if 0 <= vdist <= max_vdist and h_center_diff <= max_h_center:
                img.setdefault('captions', []).append(txt)
                consumed_text_ids.add(id(txt))
                # don't remove from elements here; we will filter consumed ids once (below)
                break

    # finally remove consumed text blocks from elements (single pass)
    elements = [el for el in elements if not (el.get('type') == 'text' and id(el) in consumed_text_ids)]
    
    elements.sort(key=get_reading_order)
    
    return elements

# ==========================================
# DOCX Writing
# ==========================================

def write_to_docx(doc, elements, page_width, page_height):
    
    # Match Page Size logic
    section = doc.sections[-1]
    # page_width/page_height are PDF points (72 pt == 1 inch).
    # Convert to inches so python-docx receives the intended size.
    section.page_width = Inches(page_width / 72.0)
    section.page_height = Inches(page_height / 72.0)
    
    # Reduce margins to allow PDF-like absolute positioning approximations
    # We use a small left margin and control the rest via indentation
    margin_buffer = 36 # 0.5 inch
    section.left_margin = Pt(margin_buffer)
    section.right_margin = Pt(margin_buffer)
    section.top_margin = Pt(margin_buffer)
    section.bottom_margin = Pt(margin_buffer)

    # --- Compute page-wide base font size once for stable heading detection ---
    all_sizes = []
    for e in elements:
        if e.get('type') != 'text':
            continue
        for ln in e.get('lines', []):
            for s in ln.get('spans', []):
                if s.get('size'):
                    all_sizes.append(float(s.get('size')))
    if all_sizes:
        # median is robust to noise (very small/very large fonts in captions)
        base_size = float(statistics.median(all_sizes))
    else:
        base_size = 12.0
        
    for el in elements:
        
        if el['type'] == 'table':
            data = el.get('data') or []
            if not data:
                continue

            # normalize ragged rows
            rows = len(data)
            cols = max((len(r) for r in data), default=0)
            cols = max(1, cols)

            table = doc.add_table(rows=rows, cols=cols)
            table.style = 'Table Grid'

            # disable autofit robustly
            try:
                table.allow_autofit = False
            except Exception:
                try:
                    table.autofit = False
                except Exception:
                    pass

            total_pts = max(1.0, el['bbox'][2] - el['bbox'][0])
            total_in = total_pts / 72.0
            col_width_in = total_in / float(cols)

            # set cell widths via tcPr (keeps prior approach)
            # --- set column widths deterministically (twips) ---
            col_twips = int(col_width_in * 1440)
            for r_idx in range(rows):
                for c_idx in range(cols):
                    cell = table.rows[r_idx].cells[c_idx]
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    # remove any existing tcW children to avoid duplicates
                    for child in list(tcPr):
                        if child.tag.endswith('tcW'):
                            tcPr.remove(child)
                    tcW = OxmlElement('w:tcW')
                    tcW.set(_qn('w:w'), str(col_twips))
                    tcW.set(_qn('w:type'), 'dxa')
                    tcPr.append(tcW)

            # fill cells, padding short rows with ""
            for r, row_data in enumerate(data):
                row_cells = table.rows[r].cells
                # ensure row_data is list-like
                row_list = list(row_data) if row_data else []
                for c in range(cols):
                    text_val = (row_list[c] if c < len(row_list) else "") or ""
                    cell = row_cells[c]
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    run = p.add_run(text_val.strip())
                    if r == 0:
                        run.bold = True
                        # header shading:
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        shd = OxmlElement('w:shd')
                        shd.set(qn('w:fill'), 'EDEDED')
                        tcPr.append(shd)
                        
            doc.add_paragraph() # spacer

        elif el['type'] == 'image':
            img_stream = io.BytesIO(el['bytes'])
            try:
                # Ensure stream position is at start
                try:
                    img_stream.seek(0)
                except Exception:
                    pass

                # Calculate width in inches based on PDF points (1/72 inch)
                img_w_pts = el['bbox'][2] - el['bbox'][0]
                img_w_inches = img_w_pts / 72.0

                # Cap to printable area (page width minus left/right margins).
                # Ensure margin_buffer is expressed in points (same units as page_width).
                printable_width_pts = max(36.0, page_width - (margin_buffer * 2))  # at least 0.5 in (36pt)
                max_img_width_in = printable_width_pts / 72.0
                img_w_inches = min(img_w_inches, max_img_width_in)

                try:
                    doc.add_picture(img_stream, width=Inches(img_w_inches))
                except Exception:
                    # fallback: add without width, let python-docx handle it
                    img_stream.seek(0)
                    doc.add_picture(img_stream)
                # paragraph containing the picture is usually the last paragraph
                last_p = doc.paragraphs[-1]

                # approximate alignment using bbox centers (unit: points)
                page_mid = page_width / 2.0
                img_mid = (el['bbox'][0] + el['bbox'][2]) / 2.0

                # threshold 20pt is fine; adjust if needed
                if abs(img_mid - page_mid) < 20:
                    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif img_mid > page_mid:
                    last_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    last_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                # After adding picture and alignment: attach captions if present
                if el.get('captions'):
                    for cap in el.get('captions'):
                        cap_text = "".join([s['text'] for l in cap.get('lines', []) for s in l.get('spans', [])]).strip()
                        if not cap_text:
                            continue

                        cap_p = doc.add_paragraph(cap_text)
                        cap_p.alignment = last_p.alignment

                        # Use 'Caption' style if available, otherwise keep default
                        caption_style_name = 'Caption'
                        if caption_style_name in (s.name for s in doc.styles):
                            try:
                                cap_p.style = caption_style_name
                            except Exception:
                                pass

                        # Try to derive a sensible font size from the caption spans
                        cap_font_size_pt = None
                        try:
                            first_line = cap.get('lines', [])[0]
                            first_span = first_line.get('spans', [])[0]
                            if first_span and first_span.get('size'):
                                cap_font_size_pt = float(first_span.get('size'))
                        except Exception:
                            cap_font_size_pt = None

                        if cap_font_size_pt is None:
                            cap_font_size_pt = 10.0  # fallback

                        # safe run update
                        try:
                            if cap_p.runs:
                                cap_p.runs[0].italic = True
                                cap_p.runs[0].font.size = Pt(max(8.0, cap_font_size_pt))
                            # spacing tweaks
                            cap_p.paragraph_format.space_before = Pt(2)
                            cap_p.paragraph_format.space_after = Pt(6)
                        except Exception:
                            pass

            except Exception as e:
                print(f"Warning: Image skipped: {e}")

        elif el['type'] == 'text':
            p = doc.add_paragraph()
            
            # Get full text to check for list patterns
            full_text = "".join([s['text'] for l in el['lines'] for s in l['spans']]).strip()
            pf = p.paragraph_format

            raw_indent = el.get('indent_pt', 0) or 0
            relative_indent = max(0.0, raw_indent - margin_buffer)

            bullet_re = r'^([\u2022\u2023\u25E6\u2043\u27A2\-\*])\s+'
            numbered_re = r'^(\d+(\.\d+)*|[A-Za-z]\.)\s+'
            # derive level more robustly (round rather than floor), avoid negative
            level = max(0, int(round(relative_indent / 18.0)))

            found_bullet = re.match(bullet_re, full_text)
            found_number = re.match(numbered_re, full_text)
            is_list = bool(found_bullet or found_number)

            if is_list:
                target_style = 'List Bullet' if found_bullet else 'List Number'
                # Prefer built-in styles if available
                if target_style in (s.name for s in doc.styles):
                    try:
                        p.style = target_style
                        pf.left_indent = Pt(12 + level * 18)
                    except Exception:
                        pf.left_indent = Pt(12 + level * 18)
                else:
                    # Emulate list: remove leading glyph from the text that will be added below
                    glyph_text = found_bullet.group(1) if found_bullet else (found_number.group(1) + '.')
                    # remove glyph from first span if present
                    first_line = el['lines'][0]
                    first_span = first_line['spans'][0]
                    if first_span and first_span.get('text', '').lstrip().startswith(glyph_text):
                        # strip the glyph from the span's text so the glyph is not duplicated
                        first_span['text'] = re.sub(r'^\s*' + re.escape(glyph_text) + r'\s*', '', first_span['text'], count=1)

                    # set hanging indent
                    pf.left_indent = Pt(18 + level * 18)
                    pf.first_line_indent = Pt(-12)
                    # Insert visible glyph run (so Word-like bullet appears)
                    glyph_run = p.add_run((glyph_text if found_bullet else '') + ' ')
                    glyph_run.bold = False

            # For nested lists, optionally set hanging indent for readability
            if p.style.name in ('List Bullet', 'List Number'):
                # hanging indent: keep bullet/number flush then text indented
                pf.first_line_indent = Pt(-12)  # negative to create hanging indent
                # Ensure a minimal left indent so bullets don't overlap margin
                if pf.left_indent and pf.left_indent.pt < 12:
                    pf.left_indent = Pt(12 + level * 12)
            
            # Smart spacing: larger gap for headings, standard gap for body text
            if p.style.name.startswith('Heading') or p.style.name == 'Title':
                pf.space_before = Pt(12)
                pf.space_after = Pt(6)
            else:
                pf.space_after = Pt(8) # Standard paragraph spacing

            # Heuristic: Calculate max font size in this paragraph
            doc_base_size = base_size  # already computed earlier in write_to_docx
            def paragraph_median_size(par_lines):
                sizes = [s.get('size') for l in par_lines for s in l.get('spans', []) if s.get('size')]
                return float(statistics.median(sizes)) if sizes else float(doc_base_size)

            p_median_size = paragraph_median_size(el['lines'])
            # proportion of spans that are bold
            all_spans = [s for l in el['lines'] for s in l.get('spans', [])]
            bold_spans = [s for s in all_spans if bool(s.get('flags', 0) & 16)]
            bold_prop = (len(bold_spans) / len(all_spans)) if all_spans else 0.0

            # Use both a relative multiplier and absolute pt delta to avoid tiny-font artifacts.
            rel_mul = p_median_size / float(doc_base_size) if doc_base_size else 1.0
            abs_delta = p_median_size - float(doc_base_size)

            heading_style = None
            styles_available = {s.name for s in doc.styles}

            # Rules (ordered):
            #  - Very large text -> Title / H1
            #  - Moderately large + boldness -> H1 / H2
            #  - Slightly larger + bold -> H2/H3
            if rel_mul >= 1.9 or p_median_size >= (doc_base_size + 8):
                heading_style = 'Title' if 'Title' in styles_available else ('Heading 1' if 'Heading 1' in styles_available else None)
            elif rel_mul >= 1.45 or p_median_size >= (doc_base_size + 4):
                heading_style = 'Heading 1' if 'Heading 1' in styles_available else ('Heading 2' if 'Heading 2' in styles_available else None)
            elif rel_mul >= 1.20 or (abs_delta >= 2 and bold_prop >= 0.35):
                heading_style = 'Heading 2' if 'Heading 2' in styles_available else ('Heading 3' if 'Heading 3' in styles_available else None)
            elif bold_prop >= 0.5 and rel_mul >= 1.05:
                heading_style = 'Heading 3' if 'Heading 3' in styles_available else None

            # If we picked a heading, apply style and modest spacing
            if heading_style:
                try:
                    p.style = heading_style
                    pf.space_before = Pt(10)
                    pf.space_after = Pt(6)
                    p.paragraph_format.keep_with_next = True
                except Exception:
                    pf.space_before = Pt(10)
                    pf.space_after = Pt(6)

            if heading_style:
                try:
                    p.style = heading_style
                    # Keep heading spacing clearer
                    pf.space_before = Pt(10)
                    pf.space_after = Pt(6)
                    p.paragraph_format.keep_with_next = True
                except Exception:
                    # fallback: no style change, but add stronger spacing
                    pf.space_before = Pt(10)
                    pf.space_after = Pt(6)
            
            for line in el['lines']:
                spans = line['spans']
                for si, span in enumerate(spans):
                    text = span.get('text') or ''
                    if not text:
                        continue

                    # If there is an explicit link in this span, keep previous hyperlink logic
                    if span.get('link'):
                        rgb = int_to_rgb(span['color'])
                        is_bold = bool(span['flags'] & 16)
                        is_italic = bool(span['flags'] & 2)
                        fname = map_font_name(span.get('font'), span.get('flags'))
                        add_hyperlink(p, span['link'], text, rgb, is_bold, is_italic, fname, span.get('size'))
                        # continue to possible spacing logic below (we still may want a space after the link)
                    else:
                        run = p.add_run(text)

                        # font name + rFonts xml
                        font_name = map_font_name(span.get('font'))
                        run.font.name = font_name
                        try:
                            rfonts = OxmlElement('w:rFonts')
                            rfonts.set(qn('w:ascii'), font_name)
                            rfonts.set(qn('w:hAnsi'), font_name)
                            rfonts.set(qn('w:eastAsia'), font_name)
                            rPr = run._element.get_or_add_rPr()
                            rPr.append(rfonts)
                        except Exception:
                            pass

                        # size
                        try:
                            if span.get('size'):
                                run.font.size = Pt(span['size'])
                        except Exception:
                            pass

                        # color
                        rgb = int_to_rgb(span.get('color'))
                        if rgb != (0, 0, 0):
                            try:
                                run.font.color.rgb = RGBColor(*rgb)
                            except Exception:
                                pass

                        # styles
                        flags = span.get('flags', 0)
                        if flags & 16:
                            run.bold = True
                        if flags & 2:
                            run.italic = True
                        if flags & 1:
                            run.font.superscript = True

                    # --- determine if we should insert a visible space between this span and the next ---
                    need_space = False
                    if si + 1 < len(spans):
                        next_span = spans[si + 1]
                        try:
                            # gap in PDF points between current span right and next span left
                            gap_pts = next_span['bbox'][0] - span['bbox'][2]
                        except Exception:
                            gap_pts = 0.0

                        # Use font size (in pts) as a scale for what counts as a visible gap
                        font_pt = span.get('size') or 12.0
                        # threshold: small minimum (1pt) plus a fraction of font size
                        gap_threshold = max(1.0, 0.18 * font_pt)

                        # Next text starts with punctuation? then usually NO space even if gap exists
                        next_text = (next_span.get('text') or '').lstrip()
                        if gap_pts > gap_threshold and not re.match(r'^[,.;:?!\)\]\}%/]', next_text):
                            need_space = True

                    if need_space:
                        # insert a space run with the same font properties so Word renders it consistently
                        space_run = p.add_run(' ')
                        try:
                            space_run.font.name = map_font_name(span.get('font'))
                            if span.get('size'):
                                space_run.font.size = Pt(span.get('size'))
                            if rgb != (0, 0, 0):
                                try:
                                    space_run.font.color.rgb = RGBColor(*rgb)
                                except Exception:
                                    pass
                            if flags & 16:
                                space_run.bold = True
                            if flags & 2:
                                space_run.italic = True
                        except Exception:
                            pass


# ==========================================
# Main
# ==========================================

def main():
    parser = argparse.ArgumentParser(description="High-Fidelity PDF to DOCX")
    parser.add_argument("input", help="Input PDF file")
    parser.add_argument("-o", "--output", help="Output DOCX file", default=None)
    args = parser.parse_args()

    input_file = args.input
    if not args.output:
        args.output = os.path.splitext(input_file)[0] + ".docx"

    if not os.path.exists(input_file):
        print(f"Error: File {input_file} not found.")
        sys.exit(1)

    print(f"Processing: {input_file}")
    
    doc_pdf = fitz.open(input_file)
    doc_word = Document()
    
    # Remove default Section to replace with PDF-sized sections
    # (Actually we just modify the existing first section)
    
    total_pages = len(doc_pdf)

    for i in range(total_pages):
        print(f"  - Converting Page {i+1}/{total_pages}...")
        page = doc_pdf[i]
        
        # Extract
        elements = extract_page_content(page)
        
        # Write
        # For multi-page, we need to ensure sections handle the breaks
        # If it's not the first page, we might need a new section for size changes
        # But for simplicity in this script, we assume uniform page sizes or just break.

        if i > 0:
            # Create a new section (force new page) so we can set page size/margins per PDF page
            doc_word.add_section(WD_SECTION.NEW_PAGE)
        
        # Pass the specific section index or let write_to_docx handle the last section

        write_to_docx(doc_word, elements, page.rect.width, page.rect.height)

    doc_word.save(args.output)
    print(f"Success! Saved to {args.output}")

if __name__ == "__main__":
    main()
