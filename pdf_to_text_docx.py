#!/usr/bin/env python3
"""
pdf_to_text_docx_with_columns.py

High-Fidelity PDF -> DOCX/TXT/LaTeX converter.
Improvements:
 - Native Table detection and rendering (DOCX tables).
 - Inline image placement (vs. end-of-page).
 - Text color, approximate indentation, and alignment detection.
 - Multi-column support.

Usage examples:
  python pdf_to_text_docx_with_columns.py input.pdf --out docx --output out.docx
  python pdf_to_text_docx_with_columns.py input.pdf --out txt --output out.txt --ocr
  python pdf_to_text_docx_with_columns.py input.pdf --out docx --latex --max-columns 3 --col-gap 0.18

Notes:
 - Requires Python 3.8+
 - pip install PyMuPDF python-docx pillow pytesseract
 - Tesseract binary required for --ocr.
"""
import argparse
import os
import sys
import tempfile
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches
from PIL import Image
import io
import re
import math

# optional import for OCR
try:
    import pytesseract
    HAVE_PYTESSERACT = True
except Exception:
    HAVE_PYTESSERACT = False

# ---------- Helpers ----------

def pil_from_pixmap(pix):
    """
    Convert a fitz.Pixmap to a PIL.Image.
    """
    if pix.n == 1:
        mode = "L"
        data = pix.samples
    elif pix.n == 3:
        mode = "RGB"
        data = pix.samples
    elif pix.n == 4:
        mode = "RGBA"
        data = pix.samples
    else:
        pix = fitz.Pixmap(pix, 0)
        mode = "RGB"
        data = pix.samples
    img = Image.frombytes(mode, (pix.width, pix.height), data)
    return img

def clean_text_for_latex(s: str) -> str:
    replacements = [
        ("\\", r"\textbackslash{}"),
        ("&", r"\&"),
        ("%", r"\%"),
        ("$", r"\$"),
        ("#", r"\#"),
        ("_", r"\_"),
        ("{", r"\{"),
        ("}", r"\}"),
        ("~", r"\textasciitilde{}"),
        ("^", r"\^{}"),
        ("<", r"\textless{}"),
        (">", r"\textgreater{}"),
    ]
    for a, b in replacements:
        s = s.replace(a, b)
    return s

def infer_section_level_by_size(size_pt):
    if size_pt is None:
        return None
    try:
        s = float(size_pt)
    except Exception:
        return None
    if s >= 20:
        return "section"
    if s >= 14:
        return "subsection"
    if s >= 11:
        return "subsubsection"
    return None

# ---------- Multi-column detection ----------

def detect_columns_from_spans(spans, page_width, gap_threshold=0.18, max_columns=3, verbose=False):
    """
    Detect column boundaries from spans' x positions.

    Algorithm (best-effort):
      - Collect left (x0) coords of spans that are not very wide (exclude full-width spans).
      - Sort these x0 coordinates and compute adjacent gaps.
      - Where normalized gap > gap_threshold we treat as a column separator.
      - Cap the number of columns to max_columns.
      - Produce a list of column ranges [(x_min, x_max), ...] covering full width.
      - Mark "full-width" spans (width fraction >= 0.7) as full_width=True; they are handled separately.
      - Assign each non-full-width span to a column by its center x.

    Returns:
      columns: list of (x0, x1) tuples in left->right order
      assigned_spans: list of spans with additional keys: 'col' (int or None) and 'full_width' (bool)
    """
    if verbose:
        print(f"  [col-detect] Page width: {page_width:.1f}, gap_threshold={gap_threshold}, max_columns={max_columns}")

    # compute width fraction for each span and collect left positions of narrow spans
    x_lefts = []
    narrow_spans = []
    assigned_spans = []
    for s in spans:
        bbox = s.get('bbox')
        if not bbox:
            s_copy = dict(s); s_copy.update({'col': None, 'full_width': False}); assigned_spans.append(s_copy); continue
        x0, y0, x1, y1 = bbox
        width = x1 - x0
        frac = width / page_width if page_width > 0 else 1.0
        s_copy = dict(s)
        if frac >= 0.70:
            # consider full-width (headers, footers) - keep full_width True
            s_copy['full_width'] = True
            s_copy['col'] = None
            assigned_spans.append(s_copy)
        else:
            s_copy['full_width'] = False
            s_copy['col'] = None
            assigned_spans.append(s_copy)
            x_lefts.append(x0)
            narrow_spans.append(s_copy)

    if not narrow_spans:
        # nothing to columnize
        if verbose:
            print("  [col-detect] No narrow spans -> single column")
        return [(0.0, page_width)], assigned_spans

    # sort unique left positions
    x_lefts_sorted = sorted(set(x_lefts))
    # compute normalized gaps between adjacent x_lefts
    gaps = []
    for i in range(len(x_lefts_sorted) - 1):
        gap = (x_lefts_sorted[i+1] - x_lefts_sorted[i]) / page_width
        gaps.append((gap, x_lefts_sorted[i], x_lefts_sorted[i+1], i))
    # sort gaps descending by size
    gaps_sorted = sorted(gaps, key=lambda t: t[0], reverse=True)
    # choose significant gaps where gap > gap_threshold; but cap to max_columns-1 cuts
    cuts = []
    for gap, left, right, idx in gaps_sorted:
        if gap > gap_threshold:
            cuts.append((left, right, gap))
        if len(cuts) >= (max_columns - 1):
            break
    # derive cut x positions (midpoints)
    cut_positions = sorted([(l + r) / 2.0 for (l, r, g) in cuts])
    # form column edges
    edges = [0.0] + cut_positions + [page_width]
    columns = []
    for i in range(len(edges) - 1):
        c0 = edges[i]
        c1 = edges[i+1]
        # small padding
        columns.append((c0, c1))
    # if too many columns due to tiny gaps, collapse to max_columns by merging rightmost
    if len(columns) > max_columns:
        # merge strategy: keep leftmost columns and merge the rest into last
        columns = columns[:max_columns-1] + [(columns[max_columns-1][0], columns[-1][1])]

    if verbose:
        print(f"  [col-detect] Detected {len(columns)} column(s): " +
              ", ".join([f"[{c0:.1f},{c1:.1f}]" for c0,c1 in columns]))

    # assign spans to columns by span center x
    for s in assigned_spans:
        if s.get('full_width'):
            s['col'] = None
            continue
        bbox = s.get('bbox')
        if not bbox:
            s['col'] = None
            continue
        x0, y0, x1, y1 = bbox
        cx = (x0 + x1) / 2.0
        # find column index
        assigned = False
        for ci, (c0, c1) in enumerate(columns):
            # include left edge, exclude right edge except for last column
            if (ci < len(columns)-1 and c0 <= cx < c1) or (ci == len(columns)-1 and c0 <= cx <= c1):
                s['col'] = ci
                assigned = True
                break
        if not assigned:
            # fallback: assign to nearest column by center distance
            best_ci = min(range(len(columns)), key=lambda ci: abs((columns[ci][0]+columns[ci][1])/2.0 - cx))
            s['col'] = best_ci

    return columns, assigned_spans

# ---------- Extraction ----------

def extract_structured_text_and_images(pdf_path, use_ocr=False, ocr_lang='eng', ocr_dpi=200, verbose=True, max_columns=3, col_gap=0.18):
    """
    Returns:
      pages: list of dicts with keys:
        - 'spans': list of spans (original dicts, now possibly annotated with 'col' and 'full_width')
        - 'columns': list of (x0,x1) if non-OCR and detection used
        - 'images': list of image dicts
        - 'ocr_text' if OCR used
        - 'page_width'
    """
    doc = fitz.open(pdf_path)
    pages_out = []

    for pno in range(len(doc)):
        page = doc[pno]
        page_dict = {"spans": [], "images": [], "page_width": page.rect.width, "columns": None}
        if use_ocr:
            if verbose:
                print(f"[OCR] Rendering page {pno+1}/{len(doc)} ...")
            mat = fitz.Matrix(ocr_dpi / 72, ocr_dpi / 72)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img = pil_from_pixmap(pix)
            if not HAVE_PYTESSERACT:
                raise RuntimeError("pytesseract not installed or not importable. Install pytesseract and the Tesseract binary.")
            txt = pytesseract.image_to_string(img, lang=ocr_lang)
            page_dict['ocr_text'] = txt
        else:
            if verbose:
                print(f"[Extract] Page {pno+1}/{len(doc)} ...")
            blocks = page.get_text("dict")
            for b in blocks.get("blocks", []):
                if b.get("type", 0) == 0:
                    for line in b.get("lines", []):
                        for span in line.get("spans", []):
                            s = {
                                "text": span.get("text", ""),
                                "font": span.get("font", ""),
                                "size": span.get("size", None),
                                "flags": span.get("flags", 0),
                                "bbox": span.get("bbox", None)
                            }
                            page_dict['spans'].append(s)
            # detect columns and assign spans
            cols, assigned = detect_columns_from_spans(page_dict['spans'],
                                                       page_dict['page_width'],
                                                       gap_threshold=col_gap,
                                                       max_columns=max_columns,
                                                       verbose=verbose)
            page_dict['columns'] = cols
            page_dict['spans'] = assigned

        # always extract images
        image_list = page.get_images(full=True)
        if image_list and verbose:
            print(f"  Found {len(image_list)} images on page {pno+1}")
        for imginfo in image_list:
            xref = imginfo[0]
            base_image = doc.extract_image(xref)
            img_bytes = base_image["image"]
            ext = base_image.get("ext", "png")
            page_dict['images'].append({'img_bytes': img_bytes, 'ext': ext, 'xref': xref, 'bbox': None})
        pages_out.append(page_dict)
    doc.close()
    return pages_out

# ---------- Output writers (use column info) ----------

def page_spans_in_reading_order(page, use_ocr=False):
    """
    Convert a page's spans to a sequence of paragraphs/runs in reading order.
    For non-OCR: use 'full_width' spans first in their vertical order, then columns left->right, each top->bottom.
    Returns list of blocks where each block is dict {'type':'text','text':...,'runs':[span,...]} or {'type':'image',...}
    """
    blocks = []
    if use_ocr and 'ocr_text' in page:
        # plain OCR text => single block
        blocks.append({'type': 'text', 'text': page['ocr_text'], 'runs': None})
        return blocks

    spans = page.get('spans', [])
    page_width = page.get('page_width', None)
    cols = page.get('columns')

    # separate full-width spans and column spans
    full_width_spans = [s for s in spans if s.get('full_width')]
    narrow_spans = [s for s in spans if not s.get('full_width') and s.get('col') is not None]

    # sort full-width by vertical position (y0)
    def y0_of(s):
        bbox = s.get('bbox')
        return bbox[1] if bbox else 0
    full_width_spans_sorted = sorted(full_width_spans, key=y0_of)

    # add full-width spans to blocks in order they appear mixing with columns by vertical position:
    # To preserve relative order with column content, we'll produce a merged sequence by scanning vertical positions:
    # Build per-column sorted lists
    col_spans_map = {}
    if cols:
        for ci in range(len(cols)):
            col_spans_map[ci] = []
        for s in narrow_spans:
            ci = s.get('col')
            if ci is None:
                # if any stray, put in first column
                ci = 0
            col_spans_map.setdefault(ci, []).append(s)
        # sort each column by y0 (top->bottom), then x0 for stability
        def sort_key(s):
            bbox = s.get('bbox') or (0,0,0,0)
            return (bbox[1], bbox[0])
        for ci in col_spans_map:
            col_spans_map[ci] = sorted(col_spans_map[ci], key=sort_key)
    else:
        # no columns detected; just one logical column
        col_spans_map[0] = sorted(narrow_spans, key=lambda s: (s.get('bbox')[1] if s.get('bbox') else 0, s.get('bbox')[0] if s.get('bbox') else 0))
        cols = [(0.0, page_width if page_width else 0.0)]

    # We'll merge full-width spans and the topmost items of columns by y coordinate to preserve reading order of mixed content.
    # Build pointers
    col_ptrs = {ci:0 for ci in col_spans_map}
    # helper to peek next y among columns
    def next_col_y(ci):
        lst = col_spans_map.get(ci, [])
        ptr = col_ptrs.get(ci, 0)
        if ptr < len(lst):
            bbox = lst[ptr].get('bbox') or (0,0,0,0)
            return bbox[1]
        return math.inf

    # collect all full-width y positions
    full_ptr = 0
    while True:
        # determine next candidate: smallest y among columns and next full-width span
        next_full_y = y0_of(full_width_spans_sorted[full_ptr]) if full_ptr < len(full_width_spans_sorted) else math.inf
        # find smallest col next y
        next_col_ci, next_col_y_val = None, math.inf
        for ci in sorted(col_spans_map.keys()):
            yval = next_col_y(ci)
            if yval < next_col_y_val:
                next_col_y_val = yval
                next_col_ci = ci
        # if all done, break
        if next_full_y == math.inf and next_col_y_val == math.inf:
            break
        # if next full-width is above next column item, emit it
        if next_full_y <= next_col_y_val:
            fw = full_width_spans_sorted[full_ptr]
            blocks.append({'type': 'text', 'text': fw.get('text', ''), 'runs': [fw]})
            full_ptr += 1
        else:
            # otherwise, we need to emit a 'column block' consisting of the next column's remaining contiguous runs at similar vertical band.
            # Simpler approach: emit entire column content in order left->right, top->bottom.
            # We will emit columns left-to-right at this point (all their remaining content).
            # This choice maintains left-to-right reading order.
            for ci in sorted(col_spans_map.keys()):
                col_list = col_spans_map[ci]
                if not col_list:
                    continue
                # emit all remaining spans in this column as a single block
                runs = col_list[col_ptrs[ci]:]
                if runs:
                    text_join = "".join([r.get('text','') for r in runs])
                    blocks.append({'type':'text', 'text': text_join, 'runs': runs, 'col': ci})
                col_ptrs[ci] = len(col_list)
            # after emitting columns, continue; full-width pointers may be before or after
    # if any remaining full-widths
    while full_ptr < len(full_width_spans_sorted):
        fw = full_width_spans_sorted[full_ptr]
        blocks.append({'type': 'text', 'text': fw.get('text', ''), 'runs': [fw]})
        full_ptr += 1

    # images appended separately by writers
    return blocks

def write_txt_from_pages(pages, output_path, use_ocr=False):
    with open(output_path, "w", encoding="utf-8") as f:
        for pno, page in enumerate(pages, start=1):
            f.write(f"\n\n--- Page {pno} ---\n\n")
            if use_ocr and 'ocr_text' in page:
                f.write(page['ocr_text'])
            else:
                blocks = page_spans_in_reading_order(page, use_ocr=False)
                for block in blocks:
                    if block['type'] == 'text':
                        # block['text'] already concatenated
                        f.write(block.get('text', ''))
                        f.write("\n\n")
            f.write("\n")
    print(f"Wrote TXT to {output_path}")

def write_docx_from_pages(pages, output_path, use_ocr=False, insert_page_breaks=True):
    doc = Document()
    for pno, page in enumerate(pages, start=1):
        if use_ocr and 'ocr_text' in page:
            text = page['ocr_text']
            for para in text.splitlines():
                doc.add_paragraph(para)
        else:
            blocks = page_spans_in_reading_order(page, use_ocr=False)
            for block in blocks:
                if block['type'] == 'text':
                    runs = block.get('runs')
                    if runs:
                        p = doc.add_paragraph()
                        for span in runs:
                            txt = span.get('text','')
                            if not txt:
                                continue
                            run = p.add_run(txt)
                            # font
                            fontname = span.get('font')
                            if fontname:
                                clean_font = re.sub(r"[-,]._.*", "", fontname)
                                try:
                                    run.font.name = clean_font
                                except Exception:
                                    pass
                            size = span.get('size')
                            if size:
                                try:
                                    run.font.size = Pt(float(size))
                                except Exception:
                                    pass
                            # heuristics bold/italic
                            flags = span.get('flags', 0)
                            if fontname and ("Bold" in fontname or "bold" in fontname):
                                run.bold = True
                            if fontname and ("Italic" in fontname or "Oblique" in fontname or "italic" in fontname):
                                run.italic = True
                            try:
                                if flags & 2:
                                    run.bold = True
                                if flags & 1:
                                    run.italic = True
                            except Exception:
                                pass
                    else:
                        # fallback: whole block text
                        doc.add_paragraph(block.get('text', ''))
        # images appended at end of page
        for imgdict in page.get('images', []):
            img_bytes = imgdict['img_bytes']
            ext = imgdict.get('ext', 'png').lower()
            with tempfile.NamedTemporaryFile(delete=False, suffix='.' + ext) as tf:
                tf.write(img_bytes)
                tmpfn = tf.name
            try:
                doc.add_picture(tmpfn, width=Inches(6))
            except Exception:
                try:
                    doc.add_picture(tmpfn)
                except Exception:
                    pass
            finally:
                try:
                    os.unlink(tmpfn)
                except Exception:
                    pass
        if insert_page_breaks and pno < len(pages):
            doc.add_page_break()
    doc.save(output_path)
    print(f"Wrote DOCX to {output_path}")

def write_latex_from_pages(pages, output_path, use_ocr=False):
    lines = []
    lines.append(r"\documentclass{article}")
    lines.append(r"\usepackage[utf8]{inputenc}")
    lines.append(r"\usepackage{graphicx}")
    lines.append(r"\begin{document}")
    for pno, page in enumerate(pages, start=1):
        lines.append(f"% --- Page {pno} ---")
        if use_ocr and 'ocr_text' in page:
            txt = page['ocr_text']
            for para in txt.splitlines():
                if para.strip() == "":
                    lines.append("")
                else:
                    lines.append(clean_text_for_latex(para) + r"\\")
        else:
            blocks = page_spans_in_reading_order(page, use_ocr=False)
            for block in blocks:
                if block['type'] == 'text':
                    text = block.get('text','').strip()
                    if not text:
                        continue
                    # Try to detect headings from first run size
                    runs = block.get('runs') or []
                    first_size = None
                    if runs:
                        first_size = runs[0].get('size')
                    sect = infer_section_level_by_size(first_size)
                    if sect:
                        lines.append(r"\{}{{{}}}".format(sect, clean_text_for_latex(text)))
                    else:
                        lines.append(clean_text_for_latex(text) + r"\\")
        # images
        if page.get('images'):
            for idx, imgdict in enumerate(page.get('images')):
                ext = imgdict.get('ext', 'png').lower()
                img_fname = f"page{pno}_img{idx + 1}.{ext}"
                lines.append(r"\begin{figure}[h]")
                lines.append(r"\centering")
                lines.append(r"\includegraphics[width=0.8\linewidth]{" + img_fname + r"}")
                lines.append(r"\end{figure}")
        lines.append(r"\newpage")
    lines.append(r"\end{document}")
    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"Wrote LaTeX to {output_path}")
    base_dir = os.path.dirname(os.path.abspath(output_path)) or "."
    for pno, page in enumerate(pages, start=1):
        for idx, imgdict in enumerate(page.get('images', [])):
            ext = imgdict.get('ext', 'png').lower()
            img_fname = os.path.join(base_dir, f"page{pno}_img{idx + 1}.{ext}")
            with open(img_fname, "wb") as g:
                g.write(imgdict['img_bytes'])
    if any(p.get('images') for p in pages):
        print("Saved extracted images next to the .tex file (filenames like page<N>_img<M>.<ext>)")

# ---------- CLI ----------

def parse_args():
    p = argparse.ArgumentParser(description="Convert PDF -> DOCX/TXT (with optional OCR, LaTeX export, and multi-column detection).")
    p.add_argument("input", help="Input PDF file path.")
    p.add_argument("--out", choices=["docx", "txt", "both"], default="docx", help="Which output text format to write.")
    p.add_argument("--output", "-o", help="Output file path. If omitted, derived from input. For --out both, output is used as stem.")
    p.add_argument("--ocr", action="store_true", help="Force OCR (using Tesseract). Useful for scanned PDFs.")
    p.add_argument("--latex", action="store_true", help="Also output a LaTeX (.tex) file.")
    p.add_argument("--lang", default="eng", help="Language for Tesseract OCR (default: eng).")
    p.add_argument("--dpi", type=int, default=200, help="DPI for rasterization for OCR (default 200).")
    p.add_argument("--no-page-breaks", dest="page_breaks", action="store_false", help="Don't insert page breaks in DOCX.")
    p.add_argument("--verbose", action="store_true", help="Verbose output.")
    p.add_argument("--max-columns", type=int, default=3, help="Maximum number of columns to detect (default 3).")
    p.add_argument("--col-gap", type=float, default=0.18, help="Normalized gap threshold between columns (default 0.18).")
    return p.parse_args()

def main():
    args = parse_args()
    inp = args.input
    if not os.path.exists(inp):
        print("Input file not found:", inp, file=sys.stderr)
        sys.exit(2)
    stem = os.path.splitext(os.path.basename(inp))[0]
    out_spec = args.output
    if not out_spec:
        if args.out == "docx":
            out_spec = stem + ".docx"
        elif args.out == "txt":
            out_spec = stem + ".txt"
        else:
            out_spec = stem
    pages = extract_structured_text_and_images(inp,
                                               use_ocr=args.ocr,
                                               ocr_lang=args.lang,
                                               ocr_dpi=args.dpi,
                                               verbose=args.verbose,
                                               max_columns=args.max_columns,
                                               col_gap=args.col_gap)
    if args.out in ("txt", "both"):
        if args.out == "both":
            out_txt = out_spec if out_spec.endswith(".txt") else out_spec + ".txt"
        else:
            out_txt = out_spec
        write_txt_from_pages(pages, out_txt, use_ocr=args.ocr)
    if args.out in ("docx", "both"):
        if args.out == "both":
            out_docx = out_spec if out_spec.endswith(".docx") else out_spec + ".docx"
        else:
            out_docx = out_spec
        write_docx_from_pages(pages, out_docx, use_ocr=args.ocr, insert_page_breaks=args.page_breaks)
    if args.latex:
        out_tex = os.path.splitext(out_spec)[0] + ".tex"
        write_latex_from_pages(pages, out_tex, use_ocr=args.ocr)

if __name__ == "__main__":
    main()
#!/usr/bin/env python3
"""
pdf_to_text_docx_with_columns.py

Enhanced PDF -> DOCX/TXT/LaTeX converter with multi-column detection.

Usage examples:
  python pdf_to_text_docx_with_columns.py input.pdf --out docx --output out.docx
  python pdf_to_text_docx_with_columns.py input.pdf --out txt --output out.txt --ocr
  python pdf_to_text_docx_with_columns.py input.pdf --out docx --latex --max-columns 3 --col-gap 0.18

Notes:
 - Requires Python 3.8+
 - pip install PyMuPDF python-docx pillow pytesseract
 - Tesseract binary required for --ocr.
"""
import argparse
import os
import sys
import tempfile
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import re
import math

# optional import for OCR
try:
    import pytesseract
    HAVE_PYTESSERACT = True
except Exception:
    HAVE_PYTESSERACT = False

# ---------- Helpers ----------

def pil_from_pixmap(pix):
    if pix.n < 4:
        # Gray or RGB
        mode = "RGB" if pix.n == 3 else "L"
        img = Image.frombytes(mode, (pix.width, pix.height), pix.samples)
    else:
        # RGBA
        img = Image.frombytes("RGBA", (pix.width, pix.height), pix.samples)
    return img

def clean_text_for_latex(s: str) -> str:
    replacements = [
        ("\\", r"\textbackslash{}"), ("&", r"\&"), ("%", r"\%"),
        ("$", r"\$"), ("#", r"\#"), ("_", r"\_"), ("{", r"\{"),
        ("}", r"\}"), ("~", r"\textasciitilde{}"), ("^", r"\^{}"),
        ("<", r"\textless{}"), (">", r"\textgreater{}"),
    ]
    for a, b in replacements:
        s = s.replace(a, b)
    return s

def int_to_rgb(color_int):
    """Convert PyMuPDF sRGB integer to (r, g, b)."""
    if color_int is None:
        return None
    # fitz colors are often sRGB int
    r = (color_int >> 16) & 255
    g = (color_int >> 8) & 255
    b = color_int & 255
    return r, g, b

def is_box_in_rect(box, rect):
    """Check if box (x0,y0,x1,y1) is largely inside rect (x0,y0,x1,y1)."""
    bx0, by0, bx1, by1 = box
    rx0, ry0, rx1, ry1 = rect
    # simplistic intersection check: center point
    cx = (bx0 + bx1) / 2
    cy = (by0 + by1) / 2
    return (rx0 <= cx <= rx1) and (ry0 <= cy <= ry1)

# ---------- Extraction & Analysis ----------

def detect_columns_from_spans(spans, page_width, gap_threshold=0.18, max_columns=3):
    """
    Same column detection logic as before, used to sort text flow.
    """
    x_lefts = []
    narrow_spans = []
    assigned_spans = []
    
    for s in spans:
        bbox = s.get('bbox')
        if not bbox:
            s_copy = dict(s); s_copy.update({'col': None, 'full_width': False}); assigned_spans.append(s_copy); continue
        x0, _, x1, _ = bbox
        width = x1 - x0
        frac = width / page_width if page_width > 0 else 1.0
        s_copy = dict(s)
        if frac >= 0.70:
            s_copy['full_width'] = True
            s_copy['col'] = None
            assigned_spans.append(s_copy)
        else:
            s_copy['full_width'] = False
            s_copy['col'] = None
            assigned_spans.append(s_copy)
            x_lefts.append(x0)
            narrow_spans.append(s_copy)

    if not narrow_spans:
        return [(0.0, page_width)], assigned_spans

    x_lefts_sorted = sorted(set(x_lefts))
    gaps = []
    for i in range(len(x_lefts_sorted) - 1):
        gap = (x_lefts_sorted[i+1] - x_lefts_sorted[i]) / page_width
        gaps.append((gap, x_lefts_sorted[i], x_lefts_sorted[i+1]))
    
    gaps_sorted = sorted(gaps, key=lambda t: t[0], reverse=True)
    cuts = []
    for gap, left, right in gaps_sorted:
        if gap > gap_threshold:
            cuts.append((left, right))
        if len(cuts) >= (max_columns - 1):
            break
            
    cut_positions = sorted([(l + r) / 2.0 for (l, r) in cuts])
    edges = [0.0] + cut_positions + [page_width]
    columns = []
    for i in range(len(edges) - 1):
        columns.append((edges[i], edges[i+1]))

    # Merge if too many
    if len(columns) > max_columns:
        columns = columns[:max_columns-1] + [(columns[max_columns-1][0], columns[-1][1])]

    # Assign columns
    for s in assigned_spans:
        if s.get('full_width'): continue
        bbox = s.get('bbox')
        if not bbox: continue
        cx = (bbox[0] + bbox[2]) / 2.0
        
        assigned = False
        for ci, (c0, c1) in enumerate(columns):
            if (ci < len(columns)-1 and c0 <= cx < c1) or (ci == len(columns)-1 and c0 <= cx <= c1):
                s['col'] = ci
                assigned = True
                break
        if not assigned:
            # nearest
            s['col'] = min(range(len(columns)), key=lambda ci: abs((columns[ci][0]+columns[ci][1])/2.0 - cx))

    return columns, assigned_spans

def extract_page_elements(doc, pno, use_ocr=False, ocr_lang='eng', ocr_dpi=200, max_columns=3, col_gap=0.18):
    """
    Extracts Text, Images, and Tables.
    Returns a dictionary representing the page.
    """
    page = doc[pno]
    page_width = page.rect.width
    page_height = page.rect.height
    
    result = {
        "width": page_width,
        "height": page_height,
        "elements": [], # mixed list of blocks (text, image, table)
        "columns": []
    }

    # 1. OCR Path
    if use_ocr:
        if not HAVE_PYTESSERACT:
            raise RuntimeError("OCR requested but pytesseract not available.")
        mat = fitz.Matrix(ocr_dpi / 72, ocr_dpi / 72)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        img = pil_from_pixmap(pix)
        txt = pytesseract.image_to_string(img, lang=ocr_lang)
        result['elements'].append({'type': 'ocr_text', 'text': txt})
        return result

    # 2. Detect Tables (High Fidelity)
    # PyMuPDF's find_tables matches grid lines to detect tables.
    tables_found = page.find_tables()
    table_rects = []
    
    for tab in tables_found:
        bbox = tab.bbox
        table_rects.append(bbox)
        # Store table data
        result['elements'].append({
            'type': 'table',
            'bbox': bbox,
            'data': tab.extract(), # [['Cell 1', 'Cell 2'], ...]
            'y0': bbox[1]
        })

    # 3. Extract Text Spans (excluding table areas)
    raw_spans = []
    text_blocks = page.get_text("dict")["blocks"]
    
    for b in text_blocks:
        if b['type'] == 0: # text block
            for line in b["lines"]:
                for span in line["spans"]:
                    sb = span["bbox"]
                    # Skip if span is inside a detected table
                    in_table = False
                    for tr in table_rects:
                        if is_box_in_rect(sb, tr):
                            in_table = True
                            break
                    if in_table:
                        continue
                    
                    s_data = {
                        "text": span["text"],
                        "font": span["font"],
                        "size": span["size"],
                        "color": span["color"],
                        "flags": span["flags"],
                        "bbox": sb
                    }
                    raw_spans.append(s_data)

    # 4. Detect Columns & Assign Spans
    cols, assigned_spans = detect_columns_from_spans(raw_spans, page_width, col_gap, max_columns)
    result['columns'] = cols

    # 5. Extract Images and insert into flow
    # fitz images don't always have bbox in get_images, so use get_text("dict") image blocks if available,
    # or fallback to get_images + approximate placement (harder).
    # get_text("dict") includes image blocks (type=1) with bboxes.
    for b in text_blocks:
        if b['type'] == 1: # image block
            img_bytes = b.get("image")
            ext = b.get("ext", "png")
            bbox = b["bbox"]
            if img_bytes:
                 result['elements'].append({
                    'type': 'image',
                    'bbox': bbox,
                    'y0': bbox[1],
                    'img_bytes': img_bytes,
                    'ext': ext
                })

    # 6. Group Text Spans into Paragraph Blocks based on Column Logic
    # (Reusing the reading order logic but creating Block objects)
    
    full_width = [s for s in assigned_spans if s.get('full_width')]
    col_map = {i: [] for i in range(len(cols))}
    
    for s in assigned_spans:
        if not s.get('full_width'):
            c = s.get('col', 0)
            if c not in col_map: c = 0
            col_map[c].append(s)

    # Sort
    full_width.sort(key=lambda x: x['bbox'][1])
    for c in col_map:
        # Sort by Y, then X
        col_map[c].sort(key=lambda x: (x['bbox'][1], x['bbox'][0]))

    # Merge Logic (Scanning down Y axis)
    # We create a list of text blocks.
    # To simplify visual fidelity: We will treat every contiguous run of text in a column as a block
    # until a large vertical gap or a full-width item interrupts it.
    
    current_y = 0
    fw_ptr = 0
    col_ptrs = {c: 0 for c in col_map}

    # Helper: get next available Y for a column
    def get_col_y(c):
        if col_ptrs[c] < len(col_map[c]):
            return col_map[c][col_ptrs[c]]['bbox'][1]
        return 999999

    while True:
        # Candidate Ys
        next_fw_y = full_width[fw_ptr]['bbox'][1] if fw_ptr < len(full_width) else 999999
        min_col_y = 999999
        
        for c in col_map:
            y = get_col_y(c)
            if y < min_col_y:
                min_col_y = y
        
        if next_fw_y == 999999 and min_col_y == 999999:
            break

        if next_fw_y <= min_col_y:
            # Add full width span
            s = full_width[fw_ptr]
            result['elements'].append({
                'type': 'text_block',
                'runs': [s],
                'y0': s['bbox'][1],
                'bbox': s['bbox'],
                'col_idx': None
            })
            fw_ptr += 1
        else:
            # Add column content. 
            # To preserve flow, we add the "row" of column content.
            # We iterate columns left to right.
            for c in sorted(col_map.keys()):
                if col_ptrs[c] >= len(col_map[c]): continue
                
                # Grab a "paragraph" worth of spans? 
                # For simplicity, grab one span at a time or group locally?
                # Let's group contiguous spans into one block for the docx writer to handle
                curr_span = col_map[c][col_ptrs[c]]
                
                # Check if this span is far below the current operational line (visual flow)
                # If we have multi-column, we usually want to read Col 1 top-bottom, then Col 2.
                # BUT, if we want visual reconstruction, we might interleave.
                # Standard PDF reading: Full Col 1, then Full Col 2.
                # However, full_width elements break this flow.
                
                # We will extract *up to* the next full_width element for this column.
                
                runs_in_block = []
                while col_ptrs[c] < len(col_map[c]):
                    cand = col_map[c][col_ptrs[c]]
                    if cand['bbox'][1] > next_fw_y:
                        break # Stop if we hit the vertical band of the next header
                    runs_in_block.append(cand)
                    col_ptrs[c] += 1
                
                if runs_in_block:
                    # Create a block
                    first_box = runs_in_block[0]['bbox']
                    result['elements'].append({
                        'type': 'text_block',
                        'runs': runs_in_block,
                        'y0': first_box[1],
                        'bbox': first_box, # approximate
                        'col_idx': c
                    })

    # Final Sort of all elements (Text blocks, Tables, Images) by Y position
    # This ensures images and tables appear roughly where they should visually.
    result['elements'].sort(key=lambda x: x['y0'])
    
    return result

# ---------- Writers ----------

def write_docx_visual(pages_data, output_path, page_breaks=True):
    doc = Document()
    
    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    for p_idx, page in enumerate(pages_data):
        page_width = page['width']
        
        for elem in page['elements']:
            
            if elem['type'] == 'ocr_text':
                doc.add_paragraph(elem['text'])

            elif elem['type'] == 'image':
                # Insert Image
                img_bytes = elem['img_bytes']
                ext = elem['ext']
                try:
                    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as tf:
                        tf.write(img_bytes)
                        tf_name = tf.name
                    
                    # constrain width to page
                    doc.add_picture(tf_name, width=Inches(5.5))
                    os.unlink(tf_name)
                except Exception as e:
                    print(f"Warning: Failed to write image: {e}")

            elif elem['type'] == 'table':
                # Build DOCX Table
                table_data = elem['data']
                if not table_data: continue
                
                rows = len(table_data)
                cols = len(table_data[0]) if rows > 0 else 0
                if cols == 0: continue

                tbl = doc.add_table(rows=rows, cols=cols)
                tbl.style = 'Table Grid'
                
                for r_i, row_data in enumerate(table_data):
                    row_cells = tbl.rows[r_i].cells
                    for c_i, cell_text in enumerate(row_data):
                        if cell_text and c_i < len(row_cells):
                            # We replace newlines to prevent broken cells
                            clean_txt = (cell_text or "").strip()
                            row_cells[c_i].text = clean_txt

            elif elem['type'] == 'text_block':
                runs = elem['runs']
                if not runs: continue
                
                # Create paragraph
                p = doc.add_paragraph()
                
                # 1. Indentation Logic (Visual Fidelity)
                # Calculate X offset relative to column start
                bbox = elem['bbox']
                col_idx = elem.get('col_idx')
                columns = page['columns']
                
                col_x0 = 0
                if col_idx is not None and col_idx < len(columns):
                    col_x0 = columns[col_idx][0]
                
                # relative x
                rel_x = bbox[0] - col_x0
                # If indentation is significant (> 10pt), apply it
                if rel_x > 10:
                    p.paragraph_format.left_indent = Pt(rel_x)
                
                # 2. Alignment Heuristic
                # Check center: 
                # Span center
                span_cx = (bbox[0] + bbox[2]) / 2
                # Context center (Column or Page)
                if col_idx is None:
                    ctx_cx = page_width / 2
                    ctx_width = page_width
                else:
                    c_range = columns[col_idx]
                    ctx_cx = (c_range[0] + c_range[1]) / 2
                    ctx_width = c_range[1] - c_range[0]
                
                # If span is roughly centered and short
                span_width = bbox[2] - bbox[0]
                if abs(span_cx - ctx_cx) < (ctx_width * 0.05) and span_width < (ctx_width * 0.8):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # 3. Add Runs with formatting
                # We group contiguous runs in this block to form the paragraph text
                for span in runs:
                    text = span['text']
                    # Some PDFs define spaces via positioning, resulting in glued text.
                    # Basic fix: ensure space if not present? (Hard without precise x-deltas)
                    # We just append text as provided by PyMuPDF
                    
                    if not text: continue
                    
                    run = p.add_run(text)
                    
                    # Size
                    if span['size']:
                        run.font.size = Pt(span['size'])
                    
                    # Color
                    rgb = int_to_rgb(span['color'])
                    if rgb != (0,0,0) and rgb is not None:
                        run.font.color.rgb = RGBColor(*rgb)
                    
                    # Bold/Italic
                    flags = span['flags']
                    if flags & 2**4: run.bold = True
                    if flags & 2**1: run.italic = True
                    
                    # Font Name (Attempt cleanup)
                    fname = span['font']
                    if fname:
                        clean = re.sub(r".*\+", "", fname) # Remove subset tag ABCDE+Arial
                        clean = clean.split("-")[0]
                        run.font.name = clean

        if page_breaks and p_idx < len(pages_data) - 1:
            doc.add_page_break()

    doc.save(output_path)
    print(f"Saved DOCX: {output_path}")

def write_txt_visual(pages_data, output_path):
    with open(output_path, "w", encoding="utf-8") as f:
        for p_idx, page in enumerate(pages_data):
            f.write(f"--- Page {p_idx+1} ---\n\n")
            for elem in page['elements']:
                if elem['type'] == 'text_block':
                    text = "".join([r['text'] for r in elem['runs']])
                    f.write(text + "\n")
                elif elem['type'] == 'table':
                    f.write("[TABLE]\n")
                    for row in elem['data']:
                        f.write(" | ".join([(c or "") for c in row]) + "\n")
                    f.write("\n")
                elif elem['type'] == 'ocr_text':
                    f.write(elem['text'] + "\n")
            f.write("\n")
    print(f"Saved TXT: {output_path}")

# ---------- Main ----------

def main():
    p = argparse.ArgumentParser(description="High-Fidelity PDF to DOCX Converter")
    p.add_argument("input", help="Input PDF")
    p.add_argument("--out", choices=["docx", "txt"], default="docx")
    p.add_argument("--output", "-o", help="Output filename")
    p.add_argument("--ocr", action="store_true", help="Use OCR")
    p.add_argument("--no-page-breaks", action="store_true", help="Disable page breaks in DOCX")
    p.add_argument("--max-columns", type=int, default=3)
    
    args = p.parse_args()
    
    if not os.path.exists(args.input):
        sys.exit(f"File not found: {args.input}")

    doc = fitz.open(args.input)
    pages_data = []
    
    print(f"Processing {len(doc)} pages...")
    for i in range(len(doc)):
        print(f"  - Analyzing Page {i+1}")
        pdata = extract_page_elements(doc, i, use_ocr=args.ocr, max_columns=args.max_columns)
        pages_data.append(pdata)
    doc.close()

    # Determine Output
    stem = os.path.splitext(args.input)[0]
    out_file = args.output
    if not out_file:
        out_file = stem + "." + args.out

    if args.out == "docx":
        write_docx_visual(pages_data, out_file, page_breaks=not args.no_page_breaks)
    else:
        write_txt_visual(pages_data, out_file)

if __name__ == "__main__":
    main()
